"""文档解析 + 父子分块。

支持 PDF / 图片(OCR) / TXT / Markdown / DOCX / XLSX。
父子文档策略：父文档 1000-1500 字（给 LLM 看），子文档 200 字 overlap 50（用于检索）。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# 支持的文件扩展名
SUPPORTED_EXTS = {".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".txt", ".md", ".docx", ".doc", ".xlsx", ".xls"}

# 章节标题正则（中文 + 英文模式）
CHAPTER_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十\d]+[章节篇部].*", re.MULTILINE),
    re.compile(r"^\d+[\.\)、]\s*\S", re.MULTILINE),       # 1. / 1) / 1、
    re.compile(r"^\d+\.\d+[\.\)、]?\s*\S", re.MULTILINE),  # 1.1 / 1.2.3
    re.compile(r"^[IVX]+[\.\)、]\s*\S", re.MULTILINE),     # I. / II.
    re.compile(r"^#+\s+\S", re.MULTILINE),                 # Markdown headings
]


def parse_file(file_path: str | Path) -> str:
    """根据扩展名路由解析器，返回提取的纯文本。"""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        return _parse_image(path)
    elif ext in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    elif ext == ".docx":
        return _parse_docx(path)
    elif ext == ".doc":
        return _parse_doc(path)
    elif ext in (".xlsx", ".xls"):
        return _parse_xlsx(path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")


def _parse_pdf(path: Path) -> str:
    """pdfplumber 逐页提取文字（流式，不一次性加载整份 PDF）。"""
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages.append(text)
            if i % 50 == 0 and i > 0:
                logger.debug("PDF 解析进度: %d 页", i + 1)
    return "\n\n".join(pages)


def _parse_image(path: Path) -> str:
    """pytesseract OCR 识别图片文字。"""
    import os
    from PIL import Image
    import pytesseract

    # 自动检测 Tesseract 安装路径
    _tesseract_paths = [
        r"D:\tesseract\tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for tp in _tesseract_paths:
        if os.path.isfile(tp):
            pytesseract.pytesseract.tesseract_cmd = tp
            break

    img = Image.open(path)
    text = pytesseract.image_to_string(img, lang="chi_sim+eng")
    return text.strip()


def _parse_docx(path: Path) -> str:
    """python-docx 提取 DOCX 文字。"""
    from docx import Document
    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n\n".join(paras)


def _parse_doc(path: Path) -> str:
    """解析旧 .doc 文件，无法解析则返回空字符串（不入库）。"""
    # 旧 .doc 是二进制格式，python-docx 偶尔能读取
    try:
        return _parse_docx(path)
    except Exception:
        pass
    # 尝试用 olefile 提取内嵌文本
    try:
        import olefile
        ole = olefile.OleFileIO(str(path))
        stream = ole.openstream('WordDocument')
        if stream:
            raw = stream.read()
            # 尝试提取可读文字
            text = raw.decode('utf-16-le', errors='ignore')
            clean = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
            if len(clean) > 50:
                return clean
    except Exception:
        pass
    return ""


def _parse_xlsx(path: Path) -> str:
    """openpyxl 提取 Excel 表格内容为结构化文本。"""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    output: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        output.append(f"## 工作表: {sheet_name}")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                output.append(" | ".join(cells))
                row_count += 1
                if row_count > 5000:
                    output.append(f"... (工作表 {sheet_name} 超过 5000 行，已截断)")
                    break
    wb.close()
    return "\n".join(output)


def cleanup_ocr(raw_text: str) -> str:
    """DeepSeek 整理 OCR 碎片化文本为通顺文档。"""
    from tools.deepseek_client import chat_json

    if len(raw_text) < 100:
        return raw_text

    prompt = f"""你是文档整理助手。以下是从图片 OCR 提取的碎片化文本，请将其整理为通顺、结构清晰的文档。

规则：
1. 修正 OCR 错误（错别字、断行、乱码）
2. 保留原始信息不增删
3. 表格数据整理为 Markdown 表格或结构化列表
4. 输出为纯文本，不要用 markdown 代码围栏

OCR 原始文本：
{raw_text[:8000]}

输出 JSON：{{"cleaned_text": "整理后的完整文本"}}"""

    try:
        result = chat_json(
            [{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=8192,
        )
        return result.get("cleaned_text", raw_text)
    except Exception as e:
        logger.warning("OCR 整理失败，使用原始文本: %s", e)
        return raw_text


# -- 文档分块 --


def split_into_parents(text: str, source_type: str = "pdf") -> list[dict]:
    """将文档拆分为父文档列表。

    每个父文档 200-1500 字，优先按章节/标题边界分割。
    过短文段自动合并，过长的进一步切分。
    """
    if not text.strip():
        return []

    # 尝试按章节标题分割
    sections = _split_by_headings(text)

    parents: list[dict] = []
    for sec in sections:
        sec_text = sec["text"].strip()
        if not sec_text:
            continue
        if len(sec_text) > 2000:
            subs = _split_long_section(sec_text, sec.get("heading", ""))
            for sub in subs:
                sub["section"] = sec.get("heading", "") or sub.get("heading", "")
                parents.append(sub)
        else:
            parents.append({
                "text": sec_text,
                "section": sec.get("heading", ""),
                "source_type": source_type,
            })

    # 合并过短的父文档（< 150 字），把短的合并到前一个
    merged: list[dict] = []
    for p in parents:
        if merged and len(p["text"]) < 150:
            merged[-1]["text"] += "\n\n" + p["text"]
            if p.get("section") and not merged[-1].get("section"):
                merged[-1]["section"] = p["section"]
        else:
            merged.append(p)

    # 过滤仍然太短的（< 30 字，无检索意义）
    merged = [m for m in merged if len(m["text"]) >= 30]

    return merged


def _split_by_headings(text: str) -> list[dict]:
    """按章节标题拆分文本。"""
    # 找到所有标题位置
    matches: list[tuple[int, int, str]] = []  # (start, end, heading_text)
    for pat in CHAPTER_PATTERNS:
        for m in pat.finditer(text):
            matches.append((m.start(), m.end(), m.group().strip()))

    if not matches:
        return [{"text": text, "heading": ""}]

    # 按位置排序
    matches.sort(key=lambda x: x[0])

    sections = []
    for i, (start, end, heading) in enumerate(matches):
        next_start = matches[i + 1][0] if i + 1 < len(matches) else len(text)
        body = text[end:next_start].strip()
        if body:
            sections.append({"text": body, "heading": heading})

    # 第一个标题之前的内容
    if matches[0][0] > 0:
        preamble = text[:matches[0][0]].strip()
        if preamble:
            sections.insert(0, {"text": preamble, "heading": ""})

    return sections


def _split_long_section(text: str, heading: str = "") -> list[dict]:
    """把太长的段落按自然段分割，合并过短的段落。"""
    paragraphs = text.split("\n\n")
    chunks: list[dict] = []
    buf = ""
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        # 累积到至少 200 字或超过上限才输出
        if buf and len(buf) + len(p) > 1500:
            chunks.append({"text": buf.strip(), "heading": heading})
            buf = p
        else:
            buf = buf + "\n\n" + p if buf else p
    if buf.strip() and len(buf.strip()) >= 30:
        chunks.append({"text": buf.strip(), "heading": heading})
    return chunks


def parent_to_children(
    parent_text: str, child_size: int = 200, overlap: int = 50
) -> list[str]:
    """将父文档切为子文档（用于检索匹配）。

    优先在句号/换行处断句，避免截断词语。
    """
    if len(parent_text) <= child_size:
        return [parent_text]

    # 按句子边界切分
    sentences = re.split(r"(?<=[。！？.!?\n])", parent_text)
    # 合并过短的句子
    merged: list[str] = []
    buf = ""
    for s in sentences:
        if not s.strip():
            continue
        if buf and len(buf) + len(s) > child_size * 2:
            merged.append(buf.strip())
            buf = s
        else:
            buf += s
    if buf.strip():
        merged.append(buf.strip())

    if not merged:
        return [parent_text]

    # 滑动窗口生成子文档
    children: list[str] = []
    step = child_size - overlap
    pos = 0
    while pos < len(merged):
        # 取足够的句子凑到 child_size
        chunk = ""
        i = pos
        while i < len(merged) and len(chunk) + len(merged[i]) < child_size + overlap:
            chunk += merged[i]
            i += 1
        if chunk.strip():
            children.append(chunk.strip())
        pos += max(1, step // max(len(merged[pos]) if pos < len(merged) else 1, 1))
        if pos >= len(merged) or len(children) > 500:
            break

    return children or [parent_text]


def _is_low_quality(text: str) -> bool:
    """检测文本是否为无检索价值的空壳内容。"""
    if len(text) < 30:
        return True  # 太短无意义
    # 检测纯路径/文件名引用
    noise_patterns = [
        "旧版 Word 文档",
        "请用 Word 打开查看",
        "文件路径:",
    ]
    for pat in noise_patterns:
        if pat in text and len(text) < 150:
            return True
    return False


def extract_metadata(
    parent_chunk: dict, file_path: str | Path, page: int = 0
) -> dict[str, Any]:
    """从父文档提取元数据。"""
    path = Path(file_path)
    return {
        "doc_title": path.stem,
        "source_file": path.name,
        "collection": "",
        "section": parent_chunk.get("section", ""),
        "page": page,
        "char_count": len(parent_chunk.get("text", "")),
        "source_type": parent_chunk.get("source_type", path.suffix.lower().lstrip(".")),
    }


# -- 完整流水线 --


def process_file(
    file_path: str | Path,
    *,
    collection: str = "",
    enable_ocr_cleanup: bool = True,
) -> list[dict]:
    """完整处理流水线：解析 → OCR整理 → 父分块 → 子分块 → 返回结构化结果。

    返回 [{
        "parent_text": "...",
        "children": ["...", "..."],
        "metadata": {...}
    }, ...]
    """
    path = Path(file_path)
    logger.info("处理文件: %s", path.name)

    # 1. 解析
    text = parse_file(path)

    # 2. 图片 OCR 整理
    ext = path.suffix.lower()
    is_image = ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp")
    if is_image and enable_ocr_cleanup and text:
        text = cleanup_ocr(text)

    if not text.strip():
        logger.warning("文件 %s 解析后无文字内容", path.name)
        return []

    # 3. 父文档分块
    parents = split_into_parents(text, source_type=ext.lstrip("."))

    # 4. 子文档切割 + 元数据（过滤空壳）
    results: list[dict] = []
    for p in parents:
        if _is_low_quality(p["text"]):
            continue
        children = parent_to_children(p["text"])
        meta = extract_metadata(p, path)
        meta["collection"] = collection
        results.append({
            "parent_text": p["text"],
            "children": children,
            "metadata": meta,
        })

    logger.info("文件 %s: %d 个父文档, %d 个子文档", path.name, len(results), sum(len(r["children"]) for r in results))
    return results
